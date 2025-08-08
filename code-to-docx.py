#!/usr/bin/env python3
"""
代码收集工具 - 用于软著申请
将指定目录下的代码文件收集到一个docx文档中
"""

import os
import sys
import json
import argparse
import fnmatch
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

class CodeCollector:
    def __init__(self):
        self.config = self._get_default_config()
        self.stats = {
            'processed_files': 0,
            'skipped_files': 0,
            'total_lines': 0,
            'skipped_reasons': [],
            'current_pages': 0,
            'skipped_by_page_limit': 0
        }
    
    def _get_default_config(self) -> Dict[str, Any]:
        """获取默认配置"""
        return {
            "input_dir": "./src",
            "output_dir": "./output",
            "filename": "project-code.docx",
            "file_extensions": [".py", ".js", ".java", ".cpp", ".c", ".h", ".html", ".css", ".ts", ".jsx", ".tsx", ".vue", ".php", ".go", ".rs", ".rb", ".swift", ".kt", ".cs", ".sql", ".xml", ".json", ".yaml", ".yml", ".md", ".txt"],
            "exclude_dirs": ["node_modules", "__pycache__", ".git", "dist", "build", "target", "bin", "obj", "vendor", ".vscode", ".idea"],
            "exclude_files": ["*.min.js", "*.min.css", "package-lock.json", "yarn.lock"],
            "max_file_size": "1MB",
            "max_pages": 0,
            "add_line_numbers": True,
            "document_title": "项目代码文档"
        }
    
    def _parse_file_size(self, size_str: str) -> int:
        """解析文件大小字符串，返回字节数"""
        if isinstance(size_str, int):
            return size_str
        
        size_str = str(size_str).upper()
        if size_str == '0':
            return 0
        
        multipliers = {
            'B': 1,
            'KB': 1024,
            'MB': 1024 * 1024,
            'GB': 1024 * 1024 * 1024
        }
        
        for unit, multiplier in multipliers.items():
            if size_str.endswith(unit):
                try:
                    number = float(size_str[:-len(unit)])
                    return int(number * multiplier)
                except ValueError:
                    pass
        
        try:
            return int(size_str)
        except ValueError:
            print(f"警告: 无法解析文件大小 '{size_str}', 使用默认值 1MB")
            return 1024 * 1024
    
    def _load_config_file(self, config_path: str) -> Dict[str, Any]:
        """从配置文件加载配置"""
        if not os.path.exists(config_path):
            return {}
        
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(f"警告: 配置文件读取失败 {config_path}: {e}")
            return {}
    
    def _parse_args(self) -> argparse.Namespace:
        """解析命令行参数"""
        parser = argparse.ArgumentParser(
            description="代码收集工具 - 将代码文件收集到docx文档中",
            formatter_class=argparse.RawDescriptionHelpFormatter,
            epilog="""
使用示例:
  python code-to-docx.py --input ./src --output ./docs
  python code-to-docx.py --config custom-config.json
  python code-to-docx.py --input /path/to/code --filename myproject.docx
            """
        )
        
        parser.add_argument('--input', '-i', help='输入目录路径')
        parser.add_argument('--output', '-o', help='输出目录路径')
        parser.add_argument('--filename', '-f', help='输出文件名 (例: code.docx)')
        parser.add_argument('--config', '-c', default='config.json', help='配置文件路径 (默认: config.json)')
        parser.add_argument('--extensions', help='文件扩展名，逗号分隔 (例: .py,.js,.java)')
        parser.add_argument('--max-size', help='最大文件大小 (例: 1MB, 500KB)')
        parser.add_argument('--max-pages', type=int, help='最大页数限制 (0表示不限制)')
        parser.add_argument('--title', help='文档标题')
        parser.add_argument('--generate-config', action='store_true', help='生成示例配置文件')
        
        return parser.parse_args()
    
    def _merge_configs(self, args: argparse.Namespace) -> Dict[str, Any]:
        """合并配置：默认配置 < 配置文件 < 命令行参数"""
        # 1. 从默认配置开始
        config = self.config.copy()
        
        # 2. 加载配置文件
        if args.config and os.path.exists(args.config):
            file_config = self._load_config_file(args.config)
            config.update(file_config)
        
        # 3. 应用命令行参数
        if args.input:
            config['input_dir'] = args.input
        if args.output:
            config['output_dir'] = args.output
        if args.filename:
            config['filename'] = args.filename
        if args.extensions:
            config['file_extensions'] = [ext.strip() for ext in args.extensions.split(',')]
        if args.max_size:
            config['max_file_size'] = args.max_size
        if args.max_pages:
            config['max_pages'] = args.max_pages
        if args.title:
            config['document_title'] = args.title
        
        return config
    
    def _validate_config(self, config: Dict[str, Any]) -> bool:
        """验证配置是否有效"""
        errors = []
        
        if not config.get('input_dir'):
            errors.append("缺少输入目录参数")
        elif not os.path.exists(config['input_dir']):
            errors.append(f"输入目录不存在: {config['input_dir']}")
        
        if not config.get('output_dir'):
            errors.append("缺少输出目录参数")
        
        if not config.get('filename'):
            errors.append("缺少输出文件名")
        elif not config['filename'].endswith('.docx'):
            errors.append("输出文件名必须以 .docx 结尾")
        
        if errors:
            print("配置错误:")
            for error in errors:
                print(f"  - {error}")
            print("\n请提供正确的参数，或使用 --help 查看帮助信息")
            return False
        
        return True
    
    def _should_skip_file(self, file_path: Path, config: Dict[str, Any]) -> Tuple[bool, Optional[str]]:
        """检查是否应该跳过文件"""
        # 检查文件扩展名
        if config['file_extensions'] and file_path.suffix.lower() not in [ext.lower() for ext in config['file_extensions']]:
            return True, f"扩展名不匹配 ({file_path.suffix})"
        
        # 检查排除的文件模式
        for pattern in config['exclude_files']:
            if fnmatch.fnmatch(file_path.name, pattern):
                return True, f"匹配排除模式 ({pattern})"
        
        # 检查文件大小
        max_size = self._parse_file_size(config['max_file_size'])
        if max_size > 0:
            try:
                file_size = file_path.stat().st_size
                if file_size > max_size:
                    size_mb = file_size / (1024 * 1024)
                    return True, f"文件过大 ({size_mb:.1f}MB)"
            except OSError:
                return True, "无法读取文件信息"
        
        return False, None
    
    def _should_skip_dir(self, dir_path: Path, config: Dict[str, Any]) -> bool:
        """检查是否应该跳过目录"""
        dir_name = dir_path.name
        return dir_name in config['exclude_dirs'] or dir_name.startswith('.')
    
    def _scan_files(self, config: Dict[str, Any]) -> List[Path]:
        """扫描目录，获取符合条件的文件列表"""
        input_dir = Path(config['input_dir'])
        files = []
        
        def scan_recursive(current_dir: Path):
            try:
                for item in current_dir.iterdir():
                    if item.is_file():
                        should_skip, reason = self._should_skip_file(item, config)
                        if not should_skip:
                            files.append(item)
                        else:
                            self.stats['skipped_files'] += 1
                            self.stats['skipped_reasons'].append(f"{item.relative_to(input_dir)}: {reason}")
                    elif item.is_dir() and not self._should_skip_dir(item, config):
                        scan_recursive(item)
            except PermissionError:
                print(f"警告: 无法访问目录 {current_dir}")
        
        scan_recursive(input_dir)
        return sorted(files)
    
    def _read_file_content(self, file_path: Path) -> Optional[str]:
        """读取文件内容，尝试多种编码"""
        encodings = ['utf-8', 'gbk', 'utf-16', 'latin1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                print(f"警告: 读取文件失败 {file_path}: {e}")
                return None
        
        print(f"警告: 无法识别文件编码 {file_path}")
        return None
    
    def _estimate_file_pages(self, file_content: str) -> int:
        """估算文件内容占用的页数"""
        if not file_content:
            return 1
        
        lines = file_content.splitlines()
        # 假设每页大约能容纳50行代码（考虑字体大小和页面边距）
        estimated_pages = max(1, (len(lines) + 49) // 50)
        return estimated_pages
    
    def _create_document(self, config: Dict[str, Any], files: List[Path]) -> Document:
        """创建docx文档"""
        doc = Document()
        max_pages = config.get('max_pages', 0)
        
        # 设置标题
        title = doc.add_heading(config['document_title'], 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加文档信息
        doc.add_paragraph(f"生成时间: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"源目录: {config['input_dir']}")
        doc.add_paragraph(f"文件总数: {len(files)}")
        
        if max_pages > 0:
            doc.add_paragraph(f"页数限制: {max_pages} 页")
        
        doc.add_page_break()
        
        # 添加目录
        doc.add_heading('文件目录', 1)
        input_dir = Path(config['input_dir'])
        
        for i, file_path in enumerate(files, 1):
            relative_path = file_path.relative_to(input_dir)
            doc.add_paragraph(f"{i}. {relative_path}")
        
        doc.add_page_break()
        
        # 初始页数：标题页 + 目录页
        self.stats['current_pages'] = 2
        
        # 添加代码内容
        for i, file_path in enumerate(files, 1):
            relative_path = file_path.relative_to(input_dir)
            
            # 读取文件内容
            content = self._read_file_content(file_path)
            if content is None:
                content = "无法读取文件内容"
            
            # 估算此文件需要的页数
            estimated_pages = self._estimate_file_pages(content)
            
            # 检查是否会超过页数限制
            if max_pages > 0 and self.stats['current_pages'] + estimated_pages > max_pages:
                # 达到页数限制，跳过剩余文件
                remaining_files = len(files) - i + 1
                self.stats['skipped_by_page_limit'] = remaining_files
                self.stats['skipped_reasons'].append(f"因页数限制跳过 {remaining_files} 个文件")
                break
            
            # 文件标题
            doc.add_heading(f"{i}. {relative_path}", 2)
            
            if content == "无法读取文件内容":
                doc.add_paragraph(content)
                self.stats['current_pages'] += 1
            else:
                # 添加代码内容
                lines = content.splitlines()
                self.stats['total_lines'] += len(lines)
                
                if config['add_line_numbers']:
                    # 带行号的代码
                    code_text = '\n'.join(f"{j+1:4d}| {line}" for j, line in enumerate(lines))
                else:
                    code_text = content
                
                # 创建代码段落，使用等宽字体
                code_para = doc.add_paragraph(code_text)
                code_para.style.font.name = 'Consolas'
                code_para.style.font.size = Inches(0.1)
                
                # 设置中文字体
                for run in code_para.runs:
                    run.font.name = 'Consolas'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            self.stats['processed_files'] += 1
            self.stats['current_pages'] += estimated_pages
            
            # 每个文件后添加分页符（除最后一个文件）
            if i < len(files):
                doc.add_page_break()
        
        return doc
    
    def _print_stats(self):
        """打印统计信息"""
        print(f"\n处理完成!")
        print(f"处理文件: {self.stats['processed_files']} 个")
        print(f"跳过文件: {self.stats['skipped_files']} 个")
        print(f"代码总行数: {self.stats['total_lines']} 行")
        print(f"文档页数: {self.stats['current_pages']} 页")
        
        if self.stats['skipped_by_page_limit'] > 0:
            print(f"因页数限制跳过文件: {self.stats['skipped_by_page_limit']} 个")
        
        if self.stats['skipped_reasons']:
            print(f"\n跳过的文件详情:")
            for reason in self.stats['skipped_reasons'][:10]:  # 只显示前10个
                print(f"  - {reason}")
            if len(self.stats['skipped_reasons']) > 10:
                print(f"  ... 还有 {len(self.stats['skipped_reasons']) - 10} 个文件被跳过")
    
    def generate_sample_config(self) -> bool:
        """生成示例配置文件"""
        config_path = "config.json"
        if os.path.exists(config_path):
            try:
                response = input(f"配置文件 {config_path} 已存在，是否覆盖? (y/N): ")
                if response.lower() != 'y':
                    print("已取消生成配置文件")
                    return False
            except (EOFError, KeyboardInterrupt):
                print("\n非交互环境，跳过覆盖确认")
                return False
        
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self._get_default_config(), f, indent=2, ensure_ascii=False)
            print(f"已生成示例配置文件: {config_path}")
            print("请编辑配置文件后重新运行脚本")
            return True
        except IOError as e:
            print(f"生成配置文件失败: {e}")
            return False
    
    def run(self) -> int:
        """运行主程序"""
        args = self._parse_args()
        
        # 生成配置文件
        if args.generate_config:
            return 0 if self.generate_sample_config() else 1
        
        # 合并配置
        config = self._merge_configs(args)
        
        # 验证配置
        if not self._validate_config(config):
            return 1
        
        # 确保输出目录存在
        os.makedirs(config['output_dir'], exist_ok=True)
        
        # 扫描文件
        print(f"正在扫描目录: {config['input_dir']}")
        files = self._scan_files(config)
        
        if not files:
            print("未找到符合条件的文件")
            return 1
        
        print(f"找到 {len(files)} 个文件，开始处理...")
        
        # 创建文档
        doc = self._create_document(config, files)
        
        # 保存文档
        output_path = os.path.join(config['output_dir'], config['filename'])
        try:
            doc.save(output_path)
            print(f"文档已保存: {output_path}")
        except Exception as e:
            print(f"保存文档失败: {e}")
            return 1
        
        # 显示统计信息
        self._print_stats()
        
        return 0


if __name__ == "__main__":
    collector = CodeCollector()
    sys.exit(collector.run())